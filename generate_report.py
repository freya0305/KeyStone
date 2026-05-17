#!/usr/bin/env python3
"""Generate KeyStone Multi-Audience Report — mature, production-grade presentation."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Color Palette ─────────────────────────────────────────────────────────────
TEAL_DARK  = RGBColor(0x00, 0x5F, 0x7D)
TEAL_LIGHT = RGBColor(0xE0, 0xF5, 0xF9)
ORANGE     = RGBColor(0xE8, 0x6A, 0x19)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x1A, 0x1A, 0x1A)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = TEAL_DARK
    return p


def body(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(space_after)
    return p


def label_value(doc, label, value, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = TEAL_DARK
    r2 = p.add_run(value)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = BLACK
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    return p


def add_page_break(doc):
    doc.add_page_break()


def styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    shade_row(hdr, "005F7D")
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        if r_idx % 2 == 0:
            shade_row(row, "F5F5F5")
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = col_widths[i]
    return table


def shade_row(row, hex_color):
    for cell in row.cells:
        set_cell_bg(cell, hex_color)


# ─── Document ─────────────────────────────────────────────────────────────────

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── COVER ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("KeyStone")
r.font.size = Pt(32); r.bold = True; r.font.color.rgb = TEAL_DARK

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Product & Technical Report")
r2.font.size = Pt(16); r2.font.color.rgb = ORANGE

doc.add_paragraph()

meta_lines = [
    ("Prepared for", "Business Manager  |  End User  |  Developer"),
    ("Version", "1.0 — May 2026"),
    ("GitHub", "github.com/freya0305/KeyStone"),
    ("Stack", "FastAPI · Next.js · PostgreSQL · Redis · Docker · Anthropic AI"),
]
for label, value in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}: "); r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = TEAL_DARK
    r2 = p.add_run(value); r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK

doc.add_paragraph()
p_note = doc.add_paragraph()
p_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = p_note.add_run(
    "KeyStone is an AI-powered job search copilot — built for Singapore's professional market, "
    "with PDPA-first compliance and production-grade infrastructure."
)
rn.italic = True; rn.font.size = Pt(10.5); rn.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

add_page_break(doc)

# ── CHAPTER 1: READING GUIDE ──────────────────────────────────────────────────
heading(doc, "1. Reading Guide")
body(doc, "This document serves three distinct audiences. The table below maps each audience to the chapters most relevant to their perspective.")

guide_data = [
    ("Business Manager",   "Chapters 2, 3, 4", "Product vision, market opportunity, competitive positioning, and financial case for launch."),
    ("End User",           "Chapters 2, 5, 6", "What KeyStone does for job seekers and recruiters, and how data is handled."),
    ("Developer",          "Chapters 5, 7",     "Compliance architecture, codebase map, testing strategy, and deployment."),
]
styled_table(doc, ["Audience", "Key Chapters", "What You Will Learn"], guide_data,
             col_widths=[Inches(1.3), Inches(1.3), Inches(3.5)])

doc.add_paragraph()
body(doc, "Chapter 2 (Product Overview) is the recommended starting point for all readers. Chapters 3 and 4 are independent and can be read in any order.", italic=True)

add_page_break(doc)

# ── CHAPTER 2: PRODUCT OVERVIEW ───────────────────────────────────────────────
heading(doc, "2. Product Overview")

heading(doc, "2.1 The Problem", 2)
body(doc, "KeyStone addresses a structural inefficiency in Singapore's job market — the absence of tooling that bridges what candidates actually offer and what employers actually need. This gap imposes costs on both sides:")

heading(doc, "For Job Seekers", 3)
seeker_pains = [
    "Every job posting requires a tailored resume. Manual rewriting consumes 45–90 minutes per application — time that scales linearly with application volume.",
    "There is no feedback loop: candidates cannot determine whether their resume content is relevant to a given role until after an rejection.",
    "Interview preparation is largely intuitive guesswork, without structured grounding in the specific job description.",
    "Application sprawl across multiple platforms (LinkedIn, JobsDB, MyCareersFuture) makes tracking status and follow-ups chaotic.",
]
for pt in seeker_pains:
    bullet(doc, pt)

heading(doc, "For Recruiters & HR Teams", 3)
recruiter_pains = [
    "Crafting accurate, compelling job descriptions takes 45 minutes to 2 hours per role — and quality varies significantly with writer experience.",
    "Without data on what distinguishes high-quality applications from poor ones, recruiters optimising JDs are working blind.",
    "Generic or poorly specified job descriptions generate high-volume, low-quality application pools — increasing screening costs for everyone.",
    "Agencies handling 50–200+ open roles per month face a scalability ceiling with manual processes.",
]
for pt in recruiter_pains:
    bullet(doc, pt)

heading(doc, "2.2 Core Features", 2)

features = [
    ("AI Resume Tailoring",
     "The candidate uploads a master resume and pastes a job description. KeyStone extracts relevant skills and experience using Haiku-class extraction, then rewrites the resume to align with the role's requirements — typically in under 30 seconds."),
    ("JD Generator",
     "The recruiter describes the role in plain language. KeyStone generates a complete, structured job description with responsibilities, requirements, and salary guidance in under 2 minutes."),
    ("Interview Copilot",
     "Grounded in the specific job description, candidates ask questions about the role, company, or positioning strategy. Responses are generated by Sonnet-class analysis, constrained to the JD context to avoid hallucination."),
    ("Application Tracker",
     "Every application is logged with status, company, role, and notes. The candidate maintains a live overview of their job search pipeline — no more spreadsheet tracking."),
    ("Secure NRIC Handling",
     "NRICs are collected only where legally required, masked before storage using a three-stage pipeline, and fully purged within 30 minutes of consent withdrawal — in compliance with PDPA."),
]

for feat_title, feat_desc in features:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{feat_title}: "); r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = TEAL_DARK
    r2 = p.add_run(feat_desc); r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK

heading(doc, "2.3 Technical Architecture", 2)
body(doc, "KeyStone is built as a containerised full-stack application. Six services run in Docker Compose:")

arch_data = [
    ("Service", "Role", "Key Technology"),
    ("web", "Next.js frontend", "Next.js 14, Tailwind CSS, shadcn/ui, Clerk auth"),
    ("api", "FastAPI backend", "FastAPI, SQLAlchemy (async), Pydantic, Alembic migrations"),
    ("worker", "Background task processor", "Celery + Redis broker for async job processing"),
    ("db", "Primary datastore", "PostgreSQL 16 with connection pooling"),
    ("cache", "Session & transient cache", "Redis 7 with AOF persistence"),
    ("localstack", "S3 dev emulator", "LocalStack S3 for resume storage (dev only)"),
]
styled_table(doc, arch_data[0], arch_data[1:],
             col_widths=[Inches(1.1), Inches(1.8), Inches(3.3)])

doc.add_paragraph()
body(doc, "AI capabilities are routed through Anthropic's API using a two-tier strategy: Haiku-class models handle extraction and parsing (high volume, low cost), while Sonnet-class models handle analysis and generation (lower volume, higher reasoning). This architecture enforces a cost ceiling of approximately SGD 5 per user per month.")

add_page_break(doc)

# ── CHAPTER 3: VALUE PROPOSITION ──────────────────────────────────────────────
heading(doc, "3. Value Proposition")

heading(doc, "3.1 Job Seekers (B2C)", 2)
b2c_rows = [
    ("Core Problem", "Each job application requires a different resume — a manually intensive, repetitive process that does not scale"),
    ("KeyStone Benefit", "Tailored resume generated in under 30 seconds, grounded in the actual job description"),
    ("Emotional Outcome", "Candidates apply with confidence, knowing their submission is specifically aligned with the role"),
    ("Time Recovery", "45–90 minutes per application returned to the job seeker"),
    ("Plan", "SGD 12 / month (Solo / Pro)"),
]
styled_table(doc, ["", ""], b2c_rows, col_widths=[Inches(2.2), Inches(4.0)])

doc.add_paragraph()
heading(doc, "3.2 Recruiters & Agencies (B2B)", 2)
b2b_rows = [
    ("Core Problem", "Job descriptions are written from scratch per role, with no data feedback loop on what attracts strong candidates"),
    ("KeyStone Benefit", "Structured JDs generated in 90 seconds, with grounded salary guidance and requirements hierarchy"),
    ("Scale", "Agencies can handle 10× more hires per month without equivalent headcount increase"),
    ("Data Compliance", "PDPA consent collection is built into the hiring workflow — no manual compliance overhead"),
    ("Agency Plan", "SGD 79–449 / month (tiered by seat volume)"),
    ("Institution Plan", "SGD 25,000–80,000 / year (university careers teams, large HR departments)"),
]
styled_table(doc, ["", ""], b2b_rows, col_widths=[Inches(2.2), Inches(4.0)])

doc.add_paragraph()
heading(doc, "3.3 Pricing", 2)
pricing_data = [
    ("Plan", "Audience", "Price", "Included"),
    ("Solo / Pro", "Job seekers", "SGD 12 / month", "AI resume tailoring, JD generator, interview copilot, application tracker"),
    ("Agency", "Recruitment agencies", "SGD 79–449 / month", "Multi-seat JD generator, candidate analytics, PDPA workflow tools"),
    ("Institution", "Universities / enterprise HR", "SGD 25K–80K / year", "Full platform, LMS integration, compliance dashboard, dedicated support"),
]
styled_table(doc, pricing_data[0], pricing_data[1:],
             col_widths=[Inches(1.1), Inches(1.4), Inches(1.5), Inches(2.2)])

add_page_break(doc)

# ── CHAPTER 4: MARKET & COMPETITION ──────────────────────────────────────────
heading(doc, "4. Market & Competition")

heading(doc, "4.1 Market Size", 2)
tam_data = [
    ("Segment", "Market Size (SGD)", "Notes"),
    ("Singapore white-collar job seekers", "~600M", "600K seekers × ~SGD 1,000 annual value per user"),
    ("APAC recruitment software", "~4.8B", "Growing 12% annually; AI integration is the primary differentiator"),
    ("Global AI resume tooling", "~18B", "SMB + enterprise; 35% CAGR through 2028"),
]
styled_table(doc, tam_data[0], tam_data[1:],
             col_widths=[Inches(2.2), Inches(1.6), Inches(2.4)])

doc.add_paragraph()
heading(doc, "4.2 Competitive Positioning", 2)
body(doc, "KeyStone's position is unique: the intersection of AI-powered resume tailoring, job description generation, and a compliance-first architecture built explicitly for Singapore's PDPA requirements. Most competitors address one or two of these; none address all three as integrated platform primitives.")
comp_data = [
    ("Competitor", "Strength", "KeyStone Advantage"),
    ("Resume.io / Novorésumé", "Strong template library", "AI tailoring + JD analysis that templates cannot provide"),
    ("LinkedIn Premium", "Professional network, brand recognition", "Tailored application content + JD generator for recruiters"),
    ("Workday / Lever", "Enterprise ATS features", "B2C simplicity and price point; built for the candidate side"),
    ("Generic job boards", "Volume and reach", "Quality and compliance over quantity; AI differentiation"),
]
styled_table(doc, comp_data[0], comp_data[1:],
             col_widths=[Inches(1.5), Inches(1.8), Inches(3.0)])

doc.add_paragraph()
heading(doc, "4.3 Risks & Mitigations", 2)
risks = [
    ("Data privacy (PDPA)",
     "NRIC triple-masking pipeline + granular six-type consent with 30-minute deletion SLA on withdrawal. Architecture is designed for PDPA audit readiness."),
    ("AI accuracy & hallucination",
     "Interview copilot responses are grounded exclusively in the provided job description context, with no general web search. Two-tier AI routing enforces a cost ceiling."),
    ("User adoption",
     "Freemium entry point reduces friction. University careers office partnerships provide word-of-mouth distribution within the primary demographic."),
    ("Regulatory change",
     "Compliance module is decoupled from core logic — PDPA amendments can be addressed without restructuring the application layer."),
]
for risk, mitigation in risks:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r1 = p.add_run(f"• {risk}: "); r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = ORANGE
    r2 = p.add_run(mitigation); r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK

add_page_break(doc)

# ── CHAPTER 5: COMPLIANCE & DATA SECURITY ────────────────────────────────────
heading(doc, "5. Compliance & Data Security")

body(doc, "KeyStone treats Singapore's Personal Data Protection Act (PDPA) as a first-class architectural constraint — not a feature added after the product was built. All data flows are designed to satisfy PDPA's collection limitation, use limitation, and protection principles.")

heading(doc, "5.1 NRIC Three-Stage Masking Pipeline", 2)
body(doc, "Singapore's NRIC number is sensitive personal data subject to strict PDPA handling requirements. KeyStone's pipeline enforces three independent stages before any NRIC data is considered stored:")
nrci_stages = [
    ("Collection Gate",
     "NRIC is collected only at the point of a specific, consent-granted transaction (e.g., phone-based signup requiring identity verification). Collection without a linked, active consent record is technically blocked."),
    ("Pre-Storage Masking",
     "Before any NRIC value is written to persistent storage, it is transformed: the first three characters are retained (citizenship class), the remaining characters are replaced with asterisks. The masked form is what reaches the database — the plaintext NRIC exists only in transient processing memory."),
    ("Withdrawal Deletion",
     "Consent withdrawal triggers a Celery background task that purges all associated NRIC data within 30 minutes. The deletion is logged for audit purposes."),
]
for stage, desc in nrci_stages:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r1 = p.add_run(f"{stage}: "); r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = TEAL_DARK
    r2 = p.add_run(desc); r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK

heading(doc, "5.2 Six-Type Consent Architecture", 2)
body(doc, "PDPA requires that consent be specific to a purpose — blanket consent is not valid. KeyStone implements six independent consent types, each tied to a distinct data use:")
consent_data = [
    ("Consent Type", "Trigger", "Data Use"),
    ("Phone signup", "User enters phone number", "Account creation and authentication"),
    ("NRIC verification", "User explicitly submits NRIC", "Identity verification where role requires it"),
    ("Resume storage", "User uploads a resume", "AI tailoring and application generation"),
    ("Application sharing", "User initiates sharing", "Sharing application data with specific recruiters"),
    ("Product analytics", "User does not opt out", "Anonymised aggregate usage data (optional; opt-out available)"),
    ("Marketing", "User opts in", "Updates on new features and pricing (user-controlled; revocable)"),
]
styled_table(doc, consent_data[0], consent_data[1:],
             col_widths=[Inches(1.6), Inches(1.8), Inches(2.8)])

heading(doc, "5.3 Security Architecture", 2)
security_items = [
    "Authentication: all API endpoints require a valid Clerk session token. Public routes are explicitly declared; all others are protected by middleware.",
    "Resume storage: AWS S3 with SSE-S3 server-side encryption. Signed, time-limited URLs replace direct object access.",
    "Secrets: Anthropic API keys, Stripe keys, and database credentials are server-side only — never exposed to the browser. INTERNAL_API_KEY protects inter-service calls within the Docker network.",
    "Network isolation: PostgreSQL and Redis run in an isolated private network segment; only the api service has direct database access.",
    "AI cost control: two-tier routing (Haiku extraction → Sonnet analysis) enforces a deterministic per-user cost ceiling, preventing runaway API spend.",
]
for item in security_items:
    bullet(doc, item)

add_page_break(doc)

# ── CHAPTER 6: USER GUIDE ─────────────────────────────────────────────────────
heading(doc, "6. User Guide")

heading(doc, "6.1 Job Seeker Experience", 2)
body(doc, "The candidate journey is designed for completion in under five minutes, end-to-end:")
seeker_steps = [
    ("Sign Up", "Register with a phone number. A granular PDPA consent form is presented — the candidate selects only the data uses they agree to. Consent is stored, timestamped, and revocable."),
    ("Upload Resume", "The candidate uploads a master resume (PDF or DOCX). KeyStone parses and structures the content as a persistent profile, ready for tailoring."),
    ("Apply", "The candidate pastes a job description from any platform. KeyStone analyses the JD and generates a tailored resume in under 30 seconds."),
    ("Interview Prep", "The candidate asks KeyStone Copilot questions about the role, the company, or their positioning strategy. Responses are grounded in the specific JD — not general web content."),
    ("Track", "Each application is logged in the tracker with status (Applied / Interview / Offer / Rejected). The candidate maintains a live overview of their pipeline."),
]
for step, desc in seeker_steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r1 = p.add_run(f"→ {step}: "); r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = TEAL_DARK
    r2 = p.add_run(desc); r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK

heading(doc, "6.2 Recruiter Experience", 2)
body(doc, "Recruiters and HR teams use KeyStone to move faster and produce higher-quality job descriptions:")
recruiter_steps = [
    ("Create JD", "The recruiter describes the role in plain language — title, seniority, industry, key requirements. KeyStone generates a complete, structured job description including responsibilities, qualifications, salary guidance, and company context."),
    ("Publish", "The generated JD is formatted and ready to post directly to any job board or ATS. The recruiter can edit any section before publishing."),
    ("Receive Applications", "Candidates who applied using KeyStone submitted resumes tailored to this exact description — meaning the quality of the applicant pool is pre-optimised."),
    ("Track Pipeline", "The recruiter dashboard provides visibility into application volume and quality indicators across all active postings."),
]
for step, desc in recruiter_steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r1 = p.add_run(f"→ {step}: "); r1.bold = True; r1.font.size = Pt(10.5); r1.font.color.rgb = ORANGE
    r2 = p.add_run(desc); r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK

add_page_break(doc)

# ── CHAPTER 7: DEVELOPER HANDOFF ──────────────────────────────────────────────
heading(doc, "7. Developer Handoff")

heading(doc, "7.1 Local Development Setup", 2)
setup_data = [
    ("Prerequisites", "Docker, Docker Compose v2+, Colima (macOS)"),
    ("Clone", "git clone https://github.com/freya0305/KeyStone"),
    ("Environment", "cp .env.example .env  # add ANTHROPIC_API_KEY, CLERK_SECRET_KEY, CLERK_PUBLISHABLE_KEY"),
    ("Start services", "docker --context colima compose up -d"),
    ("S3 bucket (first run)", "docker --context colima compose exec localstack bash -c 'awslocal s3 mb s3://keystone-resumes-dev'"),
    ("Migrations", "docker --context colima compose exec api python -m alembic upgrade head"),
]
styled_table(doc, ["Step", "Command / Action"], setup_data,
             col_widths=[Inches(1.8), Inches(4.4)])

heading(doc, "7.2 Codebase Map", 2)
locations = [
    ("Backend entry point", "keystone/main.py — FastAPI app factory, CORS, middleware, route registration"),
    ("Settings & config", "keystone/core/__init__.py — Pydantic BaseSettings; all env vars validated here"),
    ("S3 / resume storage", "keystone/services/s3.py — boto3 client, multipart upload, signed URL generation"),
    ("AI routing", "keystone/services/ai.py — two-tier Anthropic routing (Haiku extraction / Sonnet analysis), cost tracking"),
    ("NRIC pipeline", "keystone/services/nric.py — three-stage masking: collection gate, pre-storage mask, withdrawal deletion"),
    ("Consent module", "keystone/services/consent.py — six-type consent records, storage, withdrawal processing"),
    ("Stripe payments", "keystone/services/stripe.py — checkout session creation, webhook handling, subscription status"),
    ("Frontend app", "apps/web/ — Next.js 14 App Router, TypeScript, Tailwind CSS, shadcn/ui"),
    ("Clerk auth pages", "apps/web/app/(auth)/sign-in/, sign-up/ — Clerk-hosted auth flows"),
    ("JD Generator UI", "apps/web/app/(app)/new/page.tsx — recruiter-facing JD creation form"),
    ("Application tracker", "apps/web/app/(app)/resumes/page.tsx — candidate application management"),
    ("Workers", "keystone/workers/ — Celery tasks: NRIC deletion, Stripe webhooks, email notifications"),
    ("Database migrations", "alembic/ — SQLAlchemy async migration scripts, versioned with Alembic"),
    ("Docker Compose", "docker-compose.yml — six services with health checks, resource limits, logging config"),
]
styled_table(doc, ["Area", "Location / Notes"], locations,
             col_widths=[Inches(1.8), Inches(4.4)])

heading(doc, "7.3 Testing Strategy", 2)
body(doc, "KeyStone maintains a three-tier test pyramid:")
testing = [
    ("Unit tests", "pytest tests/unit/", "Fast, isolated tests for individual functions and classes. No external dependencies."),
    ("Integration tests", "pytest tests/integration/", "Uses Docker Compose test services (db, redis, localstack S3) to verify API routes and data flow end-to-end."),
    ("E2E tests", "pytest tests/e2e/", "Playwright-based browser automation covering the full candidate and recruiter flows from UI to backend."),
    ("Run all tests", "docker --context colima compose exec api pytest", ""),
]
styled_table(doc, ["Type", "Location", "Notes"], testing,
             col_widths=[Inches(1.5), Inches(1.8), Inches(3.0)])

heading(doc, "7.4 Deployment", 2)
body(doc, "Production deployments use Docker Compose with the same six-service architecture as local development. The LocalStack service is replaced with a real AWS S3 bucket. See deploy/LOCAL-DEPLOYMENT.md for the full local runbook.")
deploy_items = [
    "Provision a production AWS S3 bucket with SSE-S3 encryption enabled.",
    "Set AWS_ACCESS_KEY_ID, AWS_SECRET_ACCESS_KEY, and S3_BUCKET to production values; clear AWS_ENDPOINT_URL.",
    "Configure CLERK_SECRET_KEY and CLERK_PUBLISHABLE_KEY for production Clerk application.",
    "Set INTERNAL_API_KEY to a cryptographically random value (minimum 32 characters).",
    "Run: docker compose -f docker-compose.yml up -d --build",
    "Verify: docker compose ps — all services should report Up and healthy within their start period.",
]
for item in deploy_items:
    bullet(doc, item)

heading(doc, "7.5 Environment Variables", 2)
env_data = [
    ("Variable", "Required", "Notes"),
    ("ANTHROPIC_API_KEY", "Yes", "Anthropic API key for Haiku (extraction) and Sonnet (analysis) routing"),
    ("CLERK_SECRET_KEY", "Yes", "Clerk server-side auth; never expose to browser"),
    ("CLERK_PUBLISHABLE_KEY", "Yes", "Clerk frontend key; safe to expose in Next.js client"),
    ("STRIPE_SECRET_KEY", "Yes (payments)", "Stripe checkout and subscription management"),
    ("STRIPE_WEBHOOK_SECRET", "Yes (payments)", "Validates Stripe webhook authenticity"),
    ("AWS_ACCESS_KEY_ID", "Production", "AWS IAM user with S3 read/write for resume storage"),
    ("AWS_SECRET_ACCESS_KEY", "Production", "AWS IAM secret"),
    ("AWS_ENDPOINT_URL", "Local only", "Set to http://localstack:4566 for dev; empty in production"),
    ("S3_BUCKET", "Yes", "S3 bucket name for resume storage (keystone-resumes-dev / production)"),
    ("INTERNAL_API_KEY", "Yes", "Protects inter-service calls within the Docker network"),
    ("DATABASE_URL", "Yes", "PostgreSQL connection string (async, via asyncpg driver)"),
    ("REDIS_URL", "Yes", "Redis connection string for Celery broker and session cache"),
]
styled_table(doc, env_data[0], env_data[1:],
             col_widths=[Inches(2.0), Inches(1.2), Inches(3.0)])

heading(doc, "7.6 Pre-Launch Verification Checklist", 2)
checklist = [
    "All six Docker containers are Up and healthy (docker compose ps)",
    "Resume upload and retrieval end-to-end: upload PDF → stored in S3 → signed URL → download matches original",
    "NRIC masking verified at each stage: plaintext in processing → masked in database (S123****Z)",
    "NRIC withdrawal: consent revocation triggers Celery task → NRIC purged from database within 30 minutes",
    "PDPA consent form displayed at phone signup; all six consent types present and independently toggleable",
    "Stripe checkout: test mode payment completes → webhook fires → subscription status updated",
    "Clerk auth: sign-up, sign-in, sign-out all complete without errors; session persists across page refreshes",
    "JD Generator: plain-language input → structured, complete job description output in under 2 minutes",
    "AI Resume Tailoring: paste JD + master resume → tailored resume in under 30 seconds; output is coherent and JD-aligned",
    "Application Tracker: create application → update status → refresh page → status persists correctly",
]
for item in checklist:
    bullet(doc, item)

# ── SAVE ──────────────────────────────────────────────────────────────────────
output_path = "/Users/cell/github/project/KeyStone/KeyStone_Report.docx"
doc.save(output_path)
print(f"Saved: {output_path}")

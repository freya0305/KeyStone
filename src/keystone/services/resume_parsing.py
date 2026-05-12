"""Resume parsing service using Claude Haiku.

Handles:
- Text extraction from PDF/DOCX
- Magic byte validation
- Claude Haiku parsing for structured resume data
- SG-specific intelligence flags
"""
import hashlib
import io
import json
import re
from dataclasses import dataclass
from typing import Optional

import structlog
import pdfplumber
from docx import Document as DocxDocument

from keystone.services.nric_detector import detect_nric, assert_no_nric, mask_nric
from keystone.core import get_settings
from keystone.services.claude_client import get_claude_client, ClaudeResponse

logger = structlog.get_logger()

# Magic bytes for file validation
PDF_MAGIC = b"%PDF"
DOCX_ZIP_MAGIC = b"PK"  # DOCX is a ZIP file

# File size limits
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB


class FileValidationError(ValueError):
    """Raised when file validation fails."""
    pass


class ResumeParseError(Exception):
    """Raised when resume parsing fails."""
    pass


@dataclass
class ResumeText:
    """Extracted resume text with metadata."""
    text: str
    filename: str
    content_hash: str
    file_type: str  # "pdf" or "docx"
    page_count: Optional[int] = None
    word_count: Optional[int] = None


@dataclass
class ParsedResume:
    """Structured resume data parsed by Claude."""
    contact: dict
    summary: Optional[str]
    experience: list[dict]
    education: list[dict]
    skills: list[str]
    certifications: list[str]
    ns_status: Optional[str]  # "completed", "ongoing", "not_applicable", "unknown"
    raw_json: dict


@dataclass
class SGFlags:
    """Singapore-specific intelligence flags."""
    has_nric: bool
    has_photo: bool
    ns_quality: Optional[str]  # "full", "partial", "not_mentioned"
    ns_status: Optional[str]  # "completed", "ongoing", "not_applicable", "unknown"
    education_tier: Optional[str]  # "local_university", "polytechnic", "ite", "international", "unknown"
    pmet_signals: list[str]  # e.g., ["management_experience", "professional_title", "high_salary_expectation"]
    is_pmet: bool


def validate_file_magic_bytes(content: bytes, filename: str) -> str:
    """Validate file using magic bytes.

    Args:
        content: Raw file bytes
        filename: Original filename for extension hints

    Returns:
        File type ("pdf" or "docx")

    Raises:
        FileValidationError: If magic bytes don't match expected patterns
    """
    # Check magic bytes
    if content[: len(PDF_MAGIC)] == PDF_MAGIC:
        return "pdf"
    elif content[: len(DOCX_ZIP_MAGIC)] == DOCX_ZIP_MAGIC:
        # DOCX is a ZIP - verify it contains word/document.xml
        if b"word/document.xml" in content or _is_valid_docx(content):
            return "docx"
        raise FileValidationError("File appears to be ZIP but is not a valid DOCX")
    else:
        # Try extension-based detection as fallback
        ext = filename.lower().split(".")[-1]
        if ext == "pdf":
            return "pdf"
        elif ext in ("docx", "doc"):
            return "docx"
        raise FileValidationError(
            f"Unable to validate file type. Expected PDF or DOCX. "
            f"File may be corrupted or in an unsupported format."
        )


def _is_valid_docx(content: bytes) -> bool:
    """Check if ZIP content is a valid DOCX by looking for word/document.xml."""
    try:
        # DOCX files contain word/document.xml in the ZIP structure
        return b"word/document.xml" in content
    except Exception:
        return False


def extract_text_from_pdf(content: bytes) -> tuple[str, Optional[int]]:
    """Extract text content from PDF using pdfplumber.

    Args:
        content: Raw PDF bytes

    Returns:
        Tuple of (extracted_text, page_count)
    """
    text_parts = []
    page_count = None

    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        logger.warning("pdf_extraction_failed", error=str(e))
        raise ResumeParseError(f"Failed to extract text from PDF: {e}")

    return "\n".join(text_parts), page_count


def extract_text_from_docx(content: bytes) -> tuple[str, Optional[int]]:
    """Extract text content from DOCX using python-docx.

    Args:
        content: Raw DOCX bytes

    Returns:
        Tuple of (extracted_text, paragraph_count)
    """
    text_parts = []
    paragraph_count = None

    try:
        doc = DocxDocument(io.BytesIO(content))
        paragraph_count = len(doc.paragraphs)
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
    except Exception as e:
        logger.warning("docx_extraction_failed", error=str(e))
        raise ResumeParseError(f"Failed to extract text from DOCX: {e}")

    return "\n".join(text_parts), paragraph_count


async def extract_resume_text(
    content: bytes,
    filename: str,
) -> ResumeText:
    """Extract text from resume file with validation.

    Args:
        content: Raw file bytes
        filename: Original filename

    Returns:
        ResumeText with extracted content and metadata

    Raises:
        FileValidationError: If file is invalid
        ResumeParseError: If text extraction fails
    """
    # Validate file size
    if len(content) > MAX_FILE_SIZE:
        raise FileValidationError(
            f"File size exceeds maximum allowed size of {MAX_FILE_SIZE / 1024 / 1024}MB"
        )

    # Validate magic bytes
    file_type = validate_file_magic_bytes(content, filename)

    # Extract text based on file type
    if file_type == "pdf":
        text, page_count = extract_text_from_pdf(content)
        word_count = len(text.split()) if text else 0
    else:  # docx
        text, paragraph_count = extract_text_from_docx(content)
        word_count = len(text.split()) if text else 0
        page_count = None  # DOCX doesn't have page count easily accessible

    # Compute content hash
    content_hash = hashlib.sha256(content).hexdigest()

    logger.info(
        "resume_text_extracted",
        filename=filename,
        file_type=file_type,
        page_count=page_count,
        word_count=word_count,
        char_count=len(text),
        content_hash=content_hash[:16],
    )

    return ResumeText(
        text=text,
        filename=filename,
        content_hash=content_hash,
        file_type=file_type,
        page_count=page_count,
        word_count=word_count,
    )


def extract_sg_flags(text: str) -> SGFlags:
    """Extract Singapore-specific intelligence flags from resume text.

    Args:
        text: Resume text content

    Returns:
        SGFlags with extracted signals
    """
    text_lower = text.lower()

    # NRIC detection
    nric_result = detect_nric(text)
    has_nric = nric_result.found

    # Photo presence (heuristic)
    photo_keywords = [
        "photo attached",
        "photo included",
        "passport photo",
        "profile picture",
        "recent photo",
        "ic photo",
        "identity photo",
    ]
    has_photo = any(kw in text_lower for kw in photo_keywords)

    # NS detection
    ns_keywords_full = [
        "national service",
        "nsman",
        "nsf",
        "ns pending",
        "national serviceman",
        "saf",  # Singapore Armed Forces
        "scdf",  # Singapore Civil Defence Force
        "spf",   # Singapore Police Force
        "completed ns",
        "finished ns",
        "ns completed",
        "ns obligations",
        "2-year ns",
    ]
    ns_keywords_partial = [
        "ns",
        "national serv",
    ]

    ns_quality = "not_mentioned"
    if any(kw in text_lower for kw in ns_keywords_full):
        ns_quality = "full"
    elif any(kw in text_lower for kw in ns_keywords_partial):
        ns_quality = "partial"

    # NS status summary
    ns_status = "unknown"
    if any(kw in text_lower for kw in ["ns completed", "finished ns", "completed ns", "ns obligations fulfilled"]):
        ns_status = "completed"
    elif any(kw in text_lower for kw in ["ns ongoing", "currently serving", "ns pending", "now in ns"]):
        ns_status = "ongoing"
    elif any(kw in text_lower for kw in ["ns exemption", "exempted from ns", "not required ns"]):
        ns_status = "not_applicable"
    elif ns_quality == "not_mentioned":
        ns_status = "not_applicable"

    # Education tier detection
    education_tier = _detect_education_tier(text_lower)

    # PMET signals detection
    pmet_signals = _detect_pmet_signals(text_lower)
    is_pmet = len(pmet_signals) >= 2

    return SGFlags(
        has_nric=has_nric,
        has_photo=has_photo,
        ns_quality=ns_quality,
        ns_status=ns_status,
        education_tier=education_tier,
        pmet_signals=pmet_signals,
        is_pmet=is_pmet,
    )


def _detect_education_tier(text_lower: str) -> str:
    """Detect education tier from resume text."""
    # Local universities
    local_unis = [
        "national university of singapore", "nus", "nanyang technological university",
        "ntu", "singapore management university", "smu", "singapore university of technology",
        "sutd", "singapore institute of technology", "sit", "suss", "singapore university of social sciences",
        "sim university", "sim",
    ]
    if any(uni in text_lower for uni in local_unis):
        return "local_university"

    # Polytechnics
    polys = [
        "ngee ann polytechnic", "np", "temasek polytechnic", "tp",
        "nanyang polytechnic", "nyp", "singapore polytechnic", "sp",
        "republic polytechnic", "rp", "westminster college singapore",
    ]
    if any(poly in text_lower for poly in polys):
        return "polytechnic"

    # ITE
    if any(x in text_lower for x in [" institute of technical education", "ite ", "nitec", "higher nitec"]):
        return "ite"

    # International schools/foreign degrees
    international_keywords = [
        "university of", "college of", "bachelor", "master", "phd", "doctorate",
        "oxford", "cambridge", "imperial college", "harvard", "stanford",
        "international school", "us degree", "uk degree", "australia degree",
        "overseas university", "foreign education",
    ]
    if any(kw in text_lower for kw in international_keywords):
        # Check if it's actually a local institution vs international
        if not any(uni in text_lower for uni in local_unis):
            return "international"

    return "unknown"


def _detect_pmet_signals(text_lower: str) -> list[str]:
    """Detect PMET (Professionals, Managers, Executives, Technicians) signals."""
    signals = []

    # Management experience
    management_titles = [
        "manager", "director", "vp", "vice president", "head of", "chief",
        "lead", "supervisor", "team lead", "senior manager", "associate director",
        "general manager", "executive director",
    ]
    if any(title in text_lower for title in management_titles):
        signals.append("management_experience")

    # Professional titles
    professional_titles = [
        "engineer", "architect", "analyst", "consultant", "specialist",
        "accountant", "lawyer", "doctor", "nurse", "pharmacist", "physiotherapist",
        "teacher", "lecturer", "professor", "researcher", "scientist",
    ]
    if any(title in text_lower for title in professional_titles):
        signals.append("professional_title")

    # Salary expectations (heuristic)
    salary_keywords = ["salary expectation", "expected salary", "current salary", "notice period"]
    if any(kw in text_lower for kw in salary_keywords):
        signals.append("salary_expectation_stated")

    # High-skill industries
    high_skill_industries = [
        "banking", "finance", "technology", "it consulting", "legal", "medical",
        "pharmaceutical", "aerospace", "engineering", "data science", "ai",
        "investment banking", "management consulting",
    ]
    if any(ind in text_lower for ind in high_skill_industries):
        signals.append("high_skill_industry")

    # Degree requirements (indicates professional track)
    degree_indicators = [
        "bachelor's degree", "bachelor degree", "university graduate",
        "professional qualification", "chartered", "certified",
    ]
    if any(ind in text_lower for ind in degree_indicators):
        signals.append("degree_qualified")

    # Job hopping (career stability signal)
    if re.search(r'\d{4}\s*[-–]\s*\d{4}', text_lower) or re.search(r'[a-z]+\s*\d{4}\s*to\s*present', text_lower):
        signals.append("career_timeline_present")

    return signals


RESUME_PARSE_SYSTEM_PROMPT = """You are an expert resume parser specializing in Singapore job seekers. Extract structured information from resumes following this exact JSON schema:

{
  "contact": {
    "name": "full name or null",
    "email": "email or null",
    "phone": "phone number or null",
    "location": "singapore location or null"
  },
  "summary": "brief professional summary or null",
  "experience": [
    {
      "company": "company name",
      "title": "job title",
      "duration": "e.g., Jan 2020 - Present or Jan 2020 - Dec 2022",
      "description": "brief description of responsibilities and achievements"
    }
  ],
  "education": [
    {
      "institution": "school/university name",
      "degree": "degree or qualification",
      "year": "graduation year or null"
    }
  ],
  "skills": ["skill1", "skill2", ...],
  "certifications": ["certification1", ...] or [],
  "ns_status": "completed|ongoing|not_applicable|unknown (for Singapore National Service)"
}

Rules:
- Return ONLY valid JSON, no markdown, no explanation
- For ns_status: use "completed" if mentioned completing NS, "ongoing" if currently serving, "not_applicable" if female or exempt, "unknown" if not mentioned
- Extract as many skills as possible
- Be concise but accurate
- Missing fields should be null, not omitted
- Arrays should be empty arrays [], not null
"""


async def parse_resume_with_claude(resume_text: str, content_hash: str) -> ParsedResume:
    """Parse structured resume data using Claude Haiku.

    Args:
        resume_text: Masked resume text (NRIC already redacted)
        content_hash: SHA-256 hash for caching

    Returns:
        ParsedResume with structured data

    Raises:
        ResumeParseError: If parsing fails
    """
    # Stage 2: Assert no NRIC before sending to Claude
    assert_no_nric(resume_text)

    settings = get_settings()
    client = get_claude_client()

    try:
        response: ClaudeResponse = await client.generate(
            model=settings.anthropic_model_haiku,
            system_prompt=RESUME_PARSE_SYSTEM_PROMPT,
            user_prompt=f"Parse this resume:\n\n{resume_text[:8000]}",  # Limit to 8k chars
            timeout=10.0,  # Resume analysis ≤10s per spec
            max_tokens=2048,
        )

        # Parse JSON from response
        content = response.content.strip()
        if content.startswith("```"):
            lines = content.split("\n")
            content = "\n".join(lines[1:-1])

        raw_json = json.loads(content)

        logger.info(
            "resume_parsed",
            content_hash=content_hash[:16],
            has_contact=raw_json.get("contact", {}).get("name") is not None,
            experience_count=len(raw_json.get("experience", [])),
            skills_count=len(raw_json.get("skills", [])),
        )

        return ParsedResume(
            contact=raw_json.get("contact", {}),
            summary=raw_json.get("summary"),
            experience=raw_json.get("experience", []),
            education=raw_json.get("education", []),
            skills=raw_json.get("skills", []),
            certifications=raw_json.get("certifications", []),
            ns_status=raw_json.get("ns_status", "unknown"),
            raw_json=raw_json,
        )

    except json.JSONDecodeError as e:
        logger.warning("resume_parse_json_failed", error=str(e), content=response.content[:500] if 'response' in dir() else "N/A")
        raise ResumeParseError(f"Failed to parse resume JSON: {e}")
    except Exception as e:
        logger.error("resume_parse_failed", error=str(e))
        raise ResumeParseError(f"Failed to parse resume: {e}")


def mask_resume_text(text: str) -> str:
    """Apply NRIC Stage 1 masking to resume text.

    Args:
        text: Raw resume text

    Returns:
        Text with NRIC patterns replaced
    """
    return mask_nric(text)

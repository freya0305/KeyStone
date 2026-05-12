"""Document export service for resume generation.

Generates PDF and DOCX from resume text with accepted suggestions applied.
"""
import io
import structlog
from typing import Optional

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
except ImportError:
    A4 = None
    getSampleStyleSheet = None

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    DocxDocument = None

logger = structlog.get_logger()


def generate_pdf(context: dict) -> bytes:
    """Generate a PDF resume with suggestions applied."""
    if A4 is None:
        raise RuntimeError(
            "reportlab is required for PDF export. "
            "Install with: pip install reportlab"
        )

    buffer = __import__("io").BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    story = []

    # Title
    title_style = styles["Title"]
    story.append(Paragraph(context.get("job_title", "Resume"), title_style))
    story.append(Spacer(1, 5 * mm))

    # Company
    if context.get("company"):
        story.append(Paragraph(f"<i>{context['company']}</i>", styles["Normal"]))
        story.append(Spacer(1, 3 * mm))

    # Skills
    skills = context.get("skills", [])
    if skills:
        story.append(Paragraph("Skills: " + ", ".join(skills), styles["Normal"]))
        story.append(Spacer(1, 5 * mm))

    # Suggestions
    story.append(Paragraph("Resume Suggestions", styles["Heading2"]))
    story.append(Spacer(1, 3 * mm))

    for suggestion in context.get("suggestions", []):
        level = suggestion.get("match_level", "")
        level_color = {
            "strong": "#2e7d32",
            "transferable": "#1976d2",
            "addressable": "#f57c00",
            "fundamental": "#c62828",
        }.get(level, "#666666")

        story.append(
            Paragraph(
                f"<b>[{level.upper()}]</b> {suggestion.get('section', '')}",
                ParagraphStyle(
                    "Level",
                    textColor=__import__("reportlab.lib.colors").HexColor(level_color),
                    fontName="Helvetica-Bold",
                    fontSize=10,
                ),
            )
        )
        if suggestion.get("original_text"):
            story.append(
                Paragraph(
                    f"<b>Original:</b> {suggestion['original_text']}",
                    styles["Normal"],
                )
            )
        if suggestion.get("suggested_text"):
            story.append(
                Paragraph(
                    f"<b>Suggested:</b> {suggestion['suggested_text']}",
                    ParagraphStyle(
                        "Suggested",
                        textColor=__import__("reportlab.lib.colors").HexColor("#1b5e20"),
                        fontName="Helvetica",
                        fontSize=10,
                    ),
                )
            )
        story.append(Spacer(1, 3 * mm))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def generate_docx(context: dict) -> bytes:
    """Generate a DOCX resume with suggestions applied."""
    if DocxDocument is None:
        raise RuntimeError(
            "python-docx is required for DOCX export. "
            "Install with: pip install python-docx"
        )

    doc = DocxDocument()

    # Title
    title = doc.add_heading(context.get("job_title", "Resume"), 0)

    # Company
    if context.get("company"):
        doc.add_paragraph(context["company"]).runs[0].italic = True

    # Skills
    skills = context.get("skills", [])
    if skills:
        p = doc.add_paragraph()
        p.add_run("Skills: ").bold = True
        p.add_run(", ".join(skills))

    # Suggestions
    doc.add_heading("Resume Suggestions", level=1)

    for suggestion in context.get("suggestions", []):
        level = suggestion.get("match_level", "")
        level_colors = {
            "strong": RGBColor(0x2e, 0x7d, 0x32),
            "transferable": RGBColor(0x19, 0x76, 0xd2),
            "addressable": RGBColor(0xf5, 0x7c, 0x00),
            "fundamental": RGBColor(0xc6, 0x28, 0x28),
        }
        color = level_colors.get(level, RGBColor(0x66, 0x66, 0x66))

        p = doc.add_paragraph()
        run = p.add_run(f"[{level.upper()}] ")
        run.bold = True
        run.font.color.rgb = color
        run2 = p.add_run(suggestion.get("section", ""))
        run2.font.color.rgb = color

        if suggestion.get("original_text"):
            p2 = doc.add_paragraph()
            p2.add_run("Original: ").bold = True
            p2.add_run(suggestion["original_text"])

        if suggestion.get("suggested_text"):
            p3 = doc.add_paragraph()
            p3.add_run("Suggested: ").bold = True
            run_s = p3.add_run(suggestion["suggested_text"])
            run_s.font.color.rgb = RGBColor(0x1b, 0x5e, 0x20)

    buffer = __import__("io").BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def apply_suggestions_to_resume(
    resume_text: str,
    accepted_suggestions: list[dict],
) -> str:
    """Apply accepted suggestions to resume text by replacing original_text with suggested_text.

    Args:
        resume_text: Original resume text
        accepted_suggestions: List of dicts with original_text, suggested_text, section keys

    Returns:
        Modified resume text with suggestions applied
    """
    if not accepted_suggestions:
        return resume_text

    modified_text = resume_text
    for suggestion in accepted_suggestions:
        original = suggestion.get("original_text", "")
        suggested = suggestion.get("suggested_text", "")
        if original and suggested and original in modified_text:
            modified_text = modified_text.replace(original, suggested)
            logger.debug(
                "suggestion_applied",
                section=suggestion.get("section"),
                original_len=len(original),
                suggested_len=len(suggested),
            )

    return modified_text


async def export_resume_to_docx(
    resume_text: str,
    accepted_suggestions: list[dict],
    filename: str = "resume",
    contact_info: Optional[dict] = None,
) -> bytes:
    """Export resume to DOCX format with suggestions applied.

    Args:
        resume_text: Original resume text
        accepted_suggestions: List of accepted suggestions to apply
            Each dict should have: original_text, suggested_text, section
        filename: Output filename (without extension)
        contact_info: Optional dict with name, email, phone, location

    Returns:
        DOCX file bytes
    """
    if DocxDocument is None:
        raise RuntimeError(
            "python-docx is required for DOCX export. "
            "Install with: pip install python-docx"
        )

    # Apply accepted suggestions to resume text
    modified_text = apply_suggestions_to_resume(resume_text, accepted_suggestions)

    doc = DocxDocument()

    # Add header with contact info if provided
    if contact_info:
        if contact_info.get("name"):
            heading = doc.add_heading(contact_info["name"], 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact_parts = []
        if contact_info.get("email"):
            contact_parts.append(contact_info["email"])
        if contact_info.get("phone"):
            contact_parts.append(contact_info["phone"])
        if contact_info.get("location"):
            contact_parts.append(contact_info["location"])

        if contact_parts:
            p = doc.add_paragraph(" | ".join(contact_parts))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()

    # Parse resume text into sections and add as content
    sections = _parse_resume_sections(modified_text)

    for section_name, section_content in sections.items():
        # Add section heading
        doc.add_heading(section_name.capitalize(), level=1)

        # Add section content
        if isinstance(section_content, list):
            for item in section_content:
                p = doc.add_paragraph(item)
                p.style = doc.styles["Normal"]
        else:
            doc.add_paragraph(section_content)

        doc.add_paragraph()  # Spacing between sections

    logger.info(
        "resume_exported_docx",
        filename=filename,
        suggestions_applied=len(accepted_suggestions),
        original_len=len(resume_text),
        modified_len=len(modified_text),
    )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


async def export_resume_to_pdf(
    resume_text: str,
    accepted_suggestions: list[dict],
    filename: str = "resume",
    contact_info: Optional[dict] = None,
) -> bytes:
    """Export resume to PDF format with suggestions applied.

    Args:
        resume_text: Original resume text
        accepted_suggestions: List of accepted suggestions to apply
            Each dict should have: original_text, suggested_text, section
        filename: Output filename (without extension)
        contact_info: Optional dict with name, email, phone, location

    Returns:
        PDF file bytes
    """
    if A4 is None:
        raise RuntimeError(
            "reportlab is required for PDF export. "
            "Install with: pip install reportlab"
        )

    # Apply accepted suggestions to resume text
    modified_text = apply_suggestions_to_resume(resume_text, accepted_suggestions)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )

    styles = getSampleStyleSheet()
    story = []

    # Add header with contact info if provided
    if contact_info:
        if contact_info.get("name"):
            name_style = ParagraphStyle(
                "NameStyle",
                parent=styles["Title"],
                fontSize=18,
                alignment=1,  # CENTER
            )
            story.append(Paragraph(contact_info["name"], name_style))

        contact_parts = []
        if contact_info.get("email"):
            contact_parts.append(contact_info["email"])
        if contact_info.get("phone"):
            contact_parts.append(contact_info["phone"])
        if contact_info.get("location"):
            contact_parts.append(contact_info["location"])

        if contact_parts:
            contact_style = ParagraphStyle(
                "ContactStyle",
                parent=styles["Normal"],
                fontSize=10,
                textColor=__import__("reportlab.lib.colors").HexColor("#666666"),
                alignment=1,  # CENTER
            )
            story.append(Paragraph(" | ".join(contact_parts), contact_style))

        story.append(Spacer(1, 5 * mm))

    # Parse resume text into sections and add as content
    sections = _parse_resume_sections(modified_text)

    for section_name, section_content in sections.items():
        # Add section heading
        story.append(Paragraph(section_name.capitalize(), styles["Heading1"]))

        # Add section content
        if isinstance(section_content, list):
            for item in section_content:
                story.append(Paragraph(item, styles["Normal"]))
        else:
            story.append(Paragraph(section_content, styles["Normal"]))

        story.append(Spacer(1, 3 * mm))

    logger.info(
        "resume_exported_pdf",
        filename=filename,
        suggestions_applied=len(accepted_suggestions),
        original_len=len(resume_text),
        modified_len=len(modified_text),
    )

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def _parse_resume_sections(resume_text: str) -> dict[str, list[str]]:
    """Parse resume text into sections based on common section headers.

    Args:
        resume_text: The resume text to parse

    Returns:
        Dict mapping section names to their content (as list of paragraphs)
    """
    import re

    sections: dict[str, list[str]] = {"header": []}
    current_section = "header"
    lines = resume_text.split("\n")

    # Common section header patterns
    section_patterns = [
        r"^(summary|profile|objective|about)\s*$",
        r"^(experience|work\s+experience|employment|professional\s+experience)\s*$",
        r"^(education|academic|qualifications)\s*$",
        r"^(skills|technical\s+skills|core\s+competencies)\s*$",
        r"^(projects|personal\s+projects|portfolio)\s*$",
        r"^(certifications|certificates|licenses)\s*$",
        r"^(awards|achievements|honors)\s*$",
        r"^(volunteer|volunteering|community)\s*$",
        r"^(references|referees)\s*$",
    ]

    section_headers = {pat.strip().rstrip("$").replace("\\s*$", "").replace("\\s+", " ") for pat in section_patterns}

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Check if this line is a section header
        is_section = False
        for pattern in section_patterns:
            if re.match(pattern, line, re.IGNORECASE):
                current_section = re.sub(pattern, "", line, flags=re.IGNORECASE).strip()
                if not current_section:
                    current_section = line.lower()
                sections[current_section] = []
                is_section = True
                break

        if is_section:
            continue

        # Add line to current section
        if current_section not in sections:
            sections[current_section] = []
        sections[current_section].append(line)

    # If header section is empty, remove it
    if not sections.get("header"):
        del sections["header"]

    return sections
